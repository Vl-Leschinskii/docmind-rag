# test_chunking.py
from agents.smart_chunker import SmartChunkerAgent
from docx import Document
import traceback
import nltk
nltk.download('punkt_tab')

def test_chunking():
    print("🔍 Тестирование чанкера...")
    
    # 1. Проверяем импорт
    try:
        chunker = SmartChunkerAgent("all-MiniLM-L6-v2")
        print("✅ Чанкер создан")
    except Exception as e:
        print(f"❌ Ошибка создания чанкера: {e}")
        traceback.print_exc()
        return
    
    # 2. Проверяем на простом тексте
    test_text = "Это тестовый текст. Он состоит из нескольких предложений. Проверяем разделение на чанки."
    
    try:
        chunks = chunker.split_by_semantics(test_text)
        print(f"✅ Чанкование работает. Получено чанков: {len(chunks)}")
        for i, chunk in enumerate(chunks):
            print(f"   Чанк {i+1}: {chunk[:50]}...")
    except Exception as e:
        print(f"❌ Ошибка чанкования: {e}")
        traceback.print_exc()
    
    # 3. Проверяем документ
    try:
        doc = Document(r"docmind-rag\uploads\Анализ от Нейронаналитика.docx")
        print(f"✅ Документ открыт. Параграфов: {len(doc.paragraphs)}")
        
        # Берем первые 5 параграфов для теста
        sample_text = "\n".join([p.text for p in doc.paragraphs[:10] if p.text.strip()])
        chunks = chunker.split_by_semantics(sample_text)
        print(f"✅ Чанкование документа работает. Чанков: {len(chunks)}")
        
    except Exception as e:
        print(f"❌ Ошибка с документом: {e}")
        traceback.print_exc()

if __name__ == "__main__":
    test_chunking()