from docx import Document
import re
from typing import List, Dict, Any, Optional

class DocParserAgent:
    """Агент для парсинга Word документов и извлечения иерархии"""
    
    def __init__(self):
        self.structure = {}
    
    def parse_with_hierarchy(self, docx_path: str) -> Dict[str, Any]:
        """Извлечение структуры документа с иерархией глав/разделов"""
        try:
            doc = Document(docx_path)
        except Exception as e:
            raise Exception(f"Не удалось открыть файл {docx_path}: {e}")
        
        hierarchy = {
            'document': docx_path.split('/')[-1],
            'chapters': [],
            'total_paragraphs': len(doc.paragraphs)
        }
        
        current_chapter = None
        current_section = None
        content_buffer = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            level = self._detect_heading_level(para)
            
            if level == 1:  # Глава
                if current_chapter and content_buffer:
                    current_chapter['content'] = '\n'.join(content_buffer)
                    hierarchy['chapters'].append(current_chapter)
                
                current_chapter = {
                    'id': f"ch_{len(hierarchy['chapters']) + 1}",
                    'title': text,
                    'level': 1,
                    'sections': [],
                    'content': ''
                }
                content_buffer = []
                
            elif level == 2 and current_chapter:  # Раздел
                if current_section and content_buffer:
                    current_section['content'] = '\n'.join(content_buffer)
                    current_chapter['sections'].append(current_section)
                
                current_section = {
                    'id': f"{current_chapter['id']}_sec_{len(current_chapter['sections']) + 1}",
                    'title': text,
                    'level': 2,
                    'subsections': [],
                    'content': ''
                }
                content_buffer = []
                
            elif level >= 3 and current_section:  # Подраздел
                subsection = {
                    'id': f"{current_section['id']}_sub_{len(current_section['subsections']) + 1}",
                    'title': text,
                    'level': level,
                    'content': ''
                }
                current_section['subsections'].append(subsection)
                
            else:  # Обычный текст
                content_buffer.append(text)
        
        if current_chapter and content_buffer:
            current_chapter['content'] = '\n'.join(content_buffer)
            hierarchy['chapters'].append(current_chapter)
        
        return hierarchy
    
    def _detect_heading_level(self, paragraph) -> int:
        """Определение уровня заголовка"""
        if paragraph.style and paragraph.style.name:
            style_name = paragraph.style.name.lower()
            if 'heading' in style_name:
                try:
                    return int(style_name.replace('heading', '').strip())
                except:
                    pass
        
        text = paragraph.text.strip()
        if re.match(r'^\d+\.', text):
            return 1
        elif re.match(r'^\d+\.\d+\.', text):
            return 2
        elif re.match(r'^\d+\.\d+\.\d+\.', text):
            return 3
        
        return 0